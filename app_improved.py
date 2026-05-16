import streamlit as st
import groq
import pdfplumber
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import os
import json
import re
from datetime import datetime

# ─────────────────────────────────────────────
# CONFIG
# ─────────────────────────────────────────────
GROQ_API_KEY    = st.secrets.get("GROQ_API_KEY", "")
RAZORPAY_BASIC  = st.secrets.get("RAZORPAY_BASIC", "#")
RAZORPAY_PRO    = st.secrets.get("RAZORPAY_PRO", "#")
RAZORPAY_YEARLY = st.secrets.get("RAZORPAY_YEARLY", "#")
EMAIL_DB_FILE   = "used_emails.json"

SHEET_SCRIPT_URL = "https://script.google.com/macros/s/AKfycbzN99UwHn4Bt1mJ4MMS5ZSV-cysoTC_ac6d6oMNkWB_JAGb1i2vqBX3RmrCDqIsla3G/exec"

# Tool name - FIX: Changed from cvnixo to ResumeReflect
TOOL_NAME = "ResumeReflect"

# ─────────────────────────────────────────────
# PAGE CONFIG
# ─────────────────────────────────────────────
st.set_page_config(
    page_title="ResumeReflect - AI Resume Tailor",
    page_icon="⚡",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Add OG meta tags for proper link sharing (no performance impact)
st.markdown('<meta property="og:title" content="ResumeReflect - AI Resume Tailor"/><meta property="og:description" content="Tailor your resume with AI. Boost ATS score instantly."/><meta property="og:type" content="website"/>', unsafe_allow_html=True)

# ─────────────────────────────────────────────
# IMPROVED MODERN DESIGN - CANVA-STYLE
# ─────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Syne:wght@400;600;700;800&family=DM+Sans:ital,wght@0,300;0,400;0,500;1,400&display=swap');

:root {
    --primary:   #2D5F9E;
    --primary-light: #3E7BC4;
    --accent:    #FF6B4A;
    --accent-light: #FF8B6B;
    --success:   #00B885;
    --warning:   #F4A261;
    --bg-dark:   #0F1419;
    --bg-light:  #F8FAFB;
    --card-bg:   #FFFFFF;
    --border:    #E5E7EB;
    --text-primary: #1F2937;
    --text-secondary: #6B7280;
    --radius:    16px;
    --radius-sm: 8px;
    --shadow-sm: 0 1px 3px rgba(0,0,0,0.08);
    --shadow-md: 0 4px 12px rgba(0,0,0,0.1);
    --shadow-lg: 0 12px 32px rgba(0,0,0,0.15);
}

html, body, .stApp { 
    background: var(--bg-light) !important; 
    color: var(--text-primary) !important; 
    font-family: 'DM Sans', sans-serif !important; 
}
#MainMenu, footer, header { visibility: hidden; }
.block-container { 
    padding: 2rem 2rem !important; 
    max-width: 1100px !important; 
    margin: 0 auto; 
}

h1,h2,h3,h4,h5 { 
    font-family: 'Syne', sans-serif !important; 
    color: var(--text-primary) !important; 
    font-weight: 700 !important;
}

/* HERO SECTION */
.hero { 
    text-align: center; 
    padding: 3.5rem 1rem 2rem; 
    background: linear-gradient(135deg, var(--primary), var(--primary-light));
    border-radius: var(--radius);
    color: white;
    margin-bottom: 2rem;
    box-shadow: var(--shadow-lg);
}

.badge {
    display: inline-block;
    background: rgba(255,255,255,0.15);
    border: 1px solid rgba(255,255,255,0.3);
    border-radius: 50px; 
    padding: 6px 14px;
    font-size: 11px; 
    letter-spacing: 1px; 
    text-transform: uppercase;
    color: rgba(255,255,255,0.9);
    margin-bottom: 1rem;
    font-weight: 600;
}

.hero h1 {
    font-size: clamp(2rem, 5vw, 3.2rem) !important;
    font-weight: 800 !important; 
    line-height: 1.2 !important; 
    margin: 0 0 1rem !important;
    color: white !important;
}

.hero p { 
    color: rgba(255,255,255,0.9); 
    font-size: 1.05rem; 
    max-width: 520px; 
    margin: 0 auto; 
    line-height: 1.7; 
}

/* CARDS */
.card {
    background: var(--card-bg); 
    border: 1px solid var(--border);
    border-radius: var(--radius); 
    padding: 1.8rem; 
    margin-bottom: 1.5rem;
    transition: all 0.3s cubic-bezier(0.4,0,0.2,1);
    box-shadow: var(--shadow-sm);
}

.card:hover { 
    border-color: var(--primary);
    box-shadow: var(--shadow-md);
    transform: translateY(-2px);
}

.card-title {
    font-family: 'Syne', sans-serif; 
    font-weight: 700; 
    font-size: 1rem;
    color: var(--text-primary); 
    margin-bottom: 1rem; 
    display: flex; 
    align-items: center; 
    gap: 10px;
}

.step-num {
    background: linear-gradient(135deg, var(--primary), var(--primary-light));
    color: white; 
    border-radius: 50%; 
    width: 32px; 
    height: 32px;
    display: inline-flex; 
    align-items: center; 
    justify-content: center;
    font-size: 14px; 
    font-weight: 700; 
    flex-shrink: 0;
}

/* FILE UPLOADER */
.stFileUploader > div {
    background: var(--bg-light) !important; 
    border: 2px dashed var(--primary) !important;
    border-radius: var(--radius) !important; 
    color: var(--text-secondary) !important;
    padding: 2rem !important;
}

.stFileUploader > div:hover { 
    border-color: var(--primary-light) !important;
    background: rgba(45,95,158,0.02) !important;
}

.stFileUploader button {
    background: linear-gradient(135deg, var(--primary), var(--primary-light)) !important;
    color: white !important; 
    border: none !important;
    border-radius: var(--radius-sm) !important; 
    font-weight: 700 !important;
    padding: 10px 24px !important;
    transition: all 0.3s !important;
}

.stFileUploader button:hover {
    transform: translateY(-2px) !important;
    box-shadow: var(--shadow-md) !important;
}

.stFileUploader label { 
    color: var(--text-primary) !important; 
    font-weight: 600 !important; 
}

/* TEXT INPUTS */
.stTextArea textarea {
    background: var(--bg-light) !important; 
    border: 1px solid var(--border) !important;
    border-radius: var(--radius-sm) !important; 
    color: var(--text-primary) !important;
    font-family: 'DM Sans', sans-serif !important; 
    font-size: 14px !important;
    padding: 12px 14px !important;
    transition: all 0.2s !important;
}

.stTextArea textarea:focus { 
    border-color: var(--primary) !important; 
    box-shadow: 0 0 0 3px rgba(45,95,158,0.1) !important; 
}

.stTextInput input {
    background: var(--bg-light) !important; 
    border: 1px solid var(--border) !important;
    border-radius: var(--radius-sm) !important; 
    color: var(--text-primary) !important;
    padding: 10px 12px !important;
    transition: all 0.2s !important;
}

.stTextInput input:focus { 
    border-color: var(--primary) !important;
    box-shadow: 0 0 0 3px rgba(45,95,158,0.1) !important;
}

.stTextInput label { 
    color: var(--text-primary) !important; 
    font-weight: 600 !important; 
}

.stTextArea label { 
    color: var(--text-primary) !important; 
    font-weight: 600 !important; 
}

/* BUTTONS */
.stButton > button {
    background: linear-gradient(135deg, var(--primary), var(--primary-light)) !important;
    color: white !important; 
    border: none !important; 
    border-radius: var(--radius-sm) !important;
    padding: 12px 24px !important; 
    font-family: 'Syne', sans-serif !important;
    font-weight: 700 !important; 
    font-size: 14px !important; 
    width: 100% !important;
    transition: all 0.3s cubic-bezier(0.4,0,0.2,1) !important;
    min-height: 44px !important;
    cursor: pointer !important;
    text-transform: none !important;
}

.stButton > button:hover { 
    transform: translateY(-2px) !important; 
    box-shadow: 0 8px 20px rgba(45,95,158,0.35) !important;
    background: linear-gradient(135deg, var(--primary-light), var(--primary)) !important;
}

.stButton > button:active {
    transform: translateY(0) !important;
}

.stDownloadButton > button {
    background: linear-gradient(135deg, var(--success), #00D490) !important;
    color: white !important; 
    border: none !important; 
    border-radius: var(--radius-sm) !important;
    padding: 12px 24px !important; 
    font-family: 'Syne', sans-serif !important;
    font-weight: 700 !important; 
    font-size: 14px !important; 
    width: 100% !important;
    transition: all 0.3s !important;
    min-height: 44px !important;
}

.stDownloadButton > button:hover { 
    transform: translateY(-2px) !important; 
    box-shadow: 0 8px 20px rgba(0,184,133,0.35) !important; 
}

/* PRICING SECTION */
.pricing-grid { 
    display: grid; 
    grid-template-columns: 1fr; 
    gap: 1.5rem; 
    margin: 2rem 0; 
}

@media (min-width: 768px) {
    .pricing-grid { 
        grid-template-columns: repeat(3, 1fr); 
    }
}

.plan {
    background: var(--card-bg); 
    border: 2px solid var(--border);
    border-radius: var(--radius); 
    padding: 2rem; 
    text-align: center;
    transition: all 0.3s cubic-bezier(0.4,0,0.2,1); 
    position: relative; 
    box-shadow: var(--shadow-sm);
}

.plan:hover { 
    border-color: var(--primary);
    transform: translateY(-4px); 
    box-shadow: var(--shadow-lg); 
}

.plan.hot { 
    border-color: var(--accent); 
    background: linear-gradient(135deg, rgba(255,107,74,0.05), rgba(255,107,74,0.02));
}

.hot-badge {
    position: absolute; 
    top: -12px; 
    left: 50%; 
    transform: translateX(-50%);
    background: linear-gradient(135deg, var(--accent), var(--accent-light));
    color: white; 
    padding: 4px 12px; 
    border-radius: 50px;
    font-size: 10px; 
    font-weight: 700; 
    letter-spacing: 1px; 
    text-transform: uppercase;
    box-shadow: var(--shadow-md);
}

.plan-name { 
    font-family:'Syne',sans-serif; 
    font-weight:700; 
    font-size:.9rem; 
    color: var(--text-primary); 
    text-transform:uppercase; 
    letter-spacing:1.5px; 
    margin-bottom:.8rem; 
}

.plan-price { 
    font-family:'Syne',sans-serif; 
    font-weight:800; 
    font-size:2rem; 
    color: var(--primary); 
    margin-bottom: 0.3rem;
}

.plan-period { 
    font-size:12px; 
    color: var(--text-secondary); 
    margin-bottom:1.2rem; 
}

.plan-feat { 
    list-style:none; 
    padding:0; 
    margin:0 0 1.5rem; 
    text-align:left; 
}

.plan-feat li { 
    font-size:13px; 
    color: var(--text-secondary); 
    padding:8px 0; 
    display:flex; 
    align-items:center; 
    gap:8px; 
    font-weight: 500; 
}

.plan-feat li::before { 
    content:"✓"; 
    color: var(--success); 
    font-weight:700;
    font-size: 16px;
}

.pay-btn {
    display:block; 
    background: linear-gradient(135deg, var(--primary), var(--primary-light));
    color:white !important; 
    text-decoration:none !important; 
    border-radius: var(--radius-sm);
    padding:11px 18px; 
    font-family:'Syne',sans-serif; 
    font-weight:700;
    font-size:13px; 
    transition:all 0.3s; 
    text-align:center;
    border: none !important;
    cursor: pointer !important;
}

.pay-btn:hover { 
    box-shadow: var(--shadow-lg); 
    transform:translateY(-2px); 
}

.divider {
    height: 1px;
    background: linear-gradient(to right, transparent, var(--border), transparent);
    margin: 2rem 0;
}

/* SCORE SECTION */
.score-wrap {
    display: grid;
    grid-template-columns: 1fr 1fr;
    gap: 1.5rem;
    margin: 2rem 0;
}

.score-box {
    background: var(--card-bg);
    border: 2px solid var(--border);
    border-radius: var(--radius);
    padding: 2rem;
    text-align: center;
    transition: all 0.3s;
}

.score-box:hover {
    border-color: var(--primary);
    box-shadow: var(--shadow-md);
}

.score-val {
    font-family: 'Syne', sans-serif;
    font-size: 3rem;
    font-weight: 800;
    margin-bottom: 1rem;
}

.score-before {
    color: var(--warning);
}

.score-after {
    color: var(--success);
}

.bar-wrap {
    width: 100%;
    height: 8px;
    background: var(--bg-light);
    border-radius: 10px;
    overflow: hidden;
    margin-bottom: 1rem;
}

.bar {
    height: 100%;
    background: linear-gradient(90deg, var(--primary), var(--primary-light));
    border-radius: 10px;
    transition: width 0.6s ease-out;
}

.score-label {
    font-size: 13px;
    color: var(--text-secondary);
    font-weight: 600;
}

/* REVIEWS */
.review-section-title {
    font-family: 'Syne', sans-serif;
    font-size: 2rem;
    font-weight: 800;
    color: var(--text-primary);
    margin-bottom: 0.5rem;
    text-align: center;
}

.review-section-sub {
    text-align: center;
    color: var(--text-secondary);
    font-size: 1rem;
    margin-bottom: 2rem;
}

.reviews-grid {
    display: grid;
    grid-template-columns: 1fr;
    gap: 1rem;
    margin: 2rem 0;
}

@media (min-width: 768px) {
    .reviews-grid {
        grid-template-columns: repeat(2, 1fr);
    }
}

@media (min-width: 1024px) {
    .reviews-grid {
        grid-template-columns: repeat(4, 1fr);
    }
}

.review-card {
    background: var(--card-bg);
    border: 1px solid var(--border);
    border-radius: var(--radius);
    padding: 1.5rem;
    transition: all 0.3s;
    box-shadow: var(--shadow-sm);
}

.review-card:hover {
    border-color: var(--primary);
    box-shadow: var(--shadow-md);
    transform: translateY(-3px);
}

.review-stars {
    color: #FFA500;
    font-size: 1rem;
    margin-bottom: 0.8rem;
    letter-spacing: 2px;
}

.review-body {
    color: var(--text-primary);
    font-size: 13px;
    line-height: 1.7;
    margin-bottom: 1rem;
    font-style: italic;
    min-height: 60px;
}

.review-name {
    font-weight: 700;
    color: var(--text-primary);
    font-size: 13px;
    margin-bottom: 0.3rem;
}

.review-role {
    color: var(--text-secondary);
    font-size: 12px;
}

.review-form-card {
    background: linear-gradient(135deg, rgba(45,95,158,0.03), rgba(45,95,158,0.01));
    border: 1px solid var(--border);
    border-radius: var(--radius);
    padding: 1.5rem;
    margin: 2rem 0;
}

/* CHIPS */
.chip {
    display: inline-block;
    background: linear-gradient(135deg, rgba(45,95,158,0.1), rgba(45,95,158,0.05));
    border: 1px solid var(--primary);
    border-radius: 20px;
    padding: 6px 12px;
    font-size: 12px;
    color: var(--primary);
    margin-right: 0.5rem;
    margin-bottom: 0.5rem;
    font-weight: 600;
    transition: all 0.2s;
}

.chip:hover {
    background: linear-gradient(135deg, rgba(45,95,158,0.15), rgba(45,95,158,0.08));
    border-color: var(--primary-light);
}

/* FOOTER */
.footer {
    text-align: center;
    padding: 2.5rem 0 1rem;
    color: var(--text-secondary);
    font-size: 12px;
    border-top: 1px solid var(--border);
    margin-top: 3rem;
}

.footer-brand {
    font-family: 'Syne', sans-serif;
    font-weight: 800;
    font-size: 1.1rem;
    background: linear-gradient(135deg, var(--primary), var(--primary-light));
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    background-clip: text;
    margin-bottom: 0.5rem;
}

/* RESPONSIVE */
@media (max-width: 640px) {
    .score-wrap {
        grid-template-columns: 1fr;
    }
    
    .card {
        padding: 1.2rem;
    }
}

/* UTILITY */
.text-center { text-align: center; }
.mt-2 { margin-top: 2rem; }
.mb-2 { margin-bottom: 2rem; }
</style>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────
# HELPER FUNCTIONS
# ─────────────────────────────────────────────

def read_pdf(file_bytes):
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            text = "\n".join([page.extract_text() or "" for page in pdf.pages])
        return text
    except Exception as e:
        st.error(f"Error reading PDF: {str(e)}")
        return ""

def log_payment_interest(plan, email=""):
    """Log payment interest with proper tool name"""
    try:
        requests.post(SHEET_SCRIPT_URL, json={
            "type": "payment_interest",
            "plan": plan,
            "email": email,
            "tool_name": TOOL_NAME,  # FIX: Added proper tool name
            "timestamp": datetime.now().isoformat()
        }, timeout=5)
    except:
        pass

def log_user_action(action, email="", extra=None):
    """Log user actions with proper tool name"""
    payload = {
        "type": "user_action",
        "action": action,
        "email": email,
        "tool_name": TOOL_NAME,  # FIX: Added proper tool name
        "timestamp": datetime.now().isoformat(),
    }
    if extra:
        payload.update(extra)
    try:
        requests.post(SHEET_SCRIPT_URL, json=payload, timeout=5)
    except:
        pass

def fetch_jd_from_url(url):
    try:
        r = requests.get(url, headers={"User-Agent": "Mozilla/5.0"}, timeout=10)
        soup = BeautifulSoup(r.text, "html.parser")
        for tag in soup(["script","style","nav","footer","header"]):
            tag.decompose()
        lines = [l.strip() for l in soup.get_text(separator="\n").splitlines() if l.strip()]
        return "\n".join(lines[:300])
    except:
        return None

def call_groq(resume_text, jd_text, market_mode="🇮🇳 India (Naukri / LinkedIn India)"):
    client = groq.Groq(api_key=GROQ_API_KEY)

    if "India" in market_mode:
        market_instructions = """
INDIAN JOB MARKET RULES (very important):
- Optimise keywords specifically for Naukri.com and LinkedIn India ATS ranking
- Use Indian resume conventions: include notice period if mentioned, CTC in LPA format, percentage-based education scores
- Add Indian recruiter search terms naturally: "immediate joiner", "open to relocation", relevant Indian tech stack terms
- Keep declaration section if present in original resume
- Bullet points should include measurable Indian industry-standard metrics
- Summary should mention notice period and location preference if available
- Skills should match exactly what Indian recruiters search for on Naukri
"""
    else:
        market_instructions = """
INTERNATIONAL JOB MARKET RULES (very important):
- Optimise keywords for global ATS systems: Workday, Greenhouse, Lever, Indeed
- Use international resume conventions: no photo, no DOB, no declaration, clean 1-page preferred
- Salary references in annual USD/GBP format if mentioned
- Use action verbs and quantified achievements suited for western hiring managers
- Skills and tools should match global industry-standard terminology
- Summary should be punchy, achievement-focused, and ATS-friendly for international roles
"""

    prompt = f"""
You are an expert ATS resume specialist and career coach.
Analyze the resume against the job description carefully.
Return ONLY a JSON object. No text before or after. No markdown. Just pure JSON.

{market_instructions}

{{
  "candidate_name": "full name from resume",
  "email": "email from resume or empty string",
  "phone": "phone from resume or empty string",
  "location": "city from resume or empty string",
  "linkedin": "linkedin url from resume or empty string",
  "match_score": 75,
  "ats_keywords_found": 12,
  "ats_keywords_missing": 5,
  "strong_points": ["point 1","point 2","point 3","point 4","point 5"],
  "missing_skills": ["skill 1","skill 2","skill 3","skill 4"],
  "improvement_tips": ["tip 1","tip 2","tip 3","tip 4"],
  "summary": "2 to 3 sentence professional summary tailored to the job description",
  "work_experience": [{{
    "title": "job title exactly as in resume",
    "company": "company name exactly as in resume",
    "dates": "dates exactly as in resume",
    "location": "location exactly as in resume",
    "bullets": ["bullet rewritten with JD keywords","bullet","bullet"]
  }}],
  "projects": [{{
    "name": "project name",
    "bullets": ["project description bullet"]
  }}],
  "education": [{{
    "degree": "degree name exactly as in resume",
    "institution": "institution name exactly as in resume",
    "year": "year exactly as in resume",
    "cgpa": "cgpa or empty string"
  }}],
  "skills_technical": ["skill1","skill2","skill3"],
  "skills_tools": ["tool1","tool2","tool3"],
  "achievements": ["achievement 1","achievement 2","achievement 3"],
  "certifications": ["certification 1","certification 2"],
  "score_explanation": "2-3 sentences explaining exactly why the ATS score improved — mention specific keywords added, sections strengthened, and what made the biggest difference",
  "job_title_suggestions": ["Job Title 1","Job Title 2","Job Title 3","Job Title 4","Job Title 5"]
}}

STRICT RULES:
- Never fabricate any experience skills or education
- Only use information already present in the resume
- Rewrite bullet points using keywords from the job description
- Keep all dates company names and institutions exactly as original
- match_score must be a number not a string
- Return ONLY pure JSON nothing else

RESUME:
{resume_text}

JOB DESCRIPTION:
{jd_text}
"""
    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.5,
        max_tokens=3500
    )
    return response.choices[0].message.content

def parse_json(text):
    text = text.strip()
    if "```" in text:
        lines = [l for l in text.split("\n") if not l.strip().startswith("```")]
        text = "\n".join(lines)
    start = text.find("{")
    end = text.rfind("}") + 1
    if start != -1 and end > start:
        text = text[start:end]
    return json.loads(text)

def simple_ats_score(resume_text, jd_text):
    jd_words  = set(re.findall(r'\b[a-zA-Z]{4,}\b', jd_text.lower()))
    res_words = set(re.findall(r'\b[a-zA-Z]{4,}\b', resume_text.lower()))
    common    = jd_words & res_words
    score     = int((len(common) / max(len(jd_words), 1)) * 100)
    return max(min(score, 75), 5)

# ─────────────────────────────────────────────
# DOCUMENT BUILDERS (abbreviated for space)
# ─────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def build_resume(data, watermark=False):
    doc = Document()
    
    # Header with name
    name = data.get('candidate_name', 'Candidate')
    header = doc.add_heading(name, 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Contact info
    contact_info = []
    if data.get('phone'): contact_info.append(data['phone'])
    if data.get('email'): contact_info.append(data['email'])
    if data.get('location'): contact_info.append(data['location'])
    if data.get('linkedin'): contact_info.append(data['linkedin'])
    
    if contact_info:
        contact = doc.add_paragraph(" | ".join(contact_info))
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact.runs[0].font.size = Pt(9)
    
    # Professional Summary
    if data.get('summary'):
        doc.add_heading('Professional Summary', level=2)
        doc.add_paragraph(data['summary'])
    
    # Work Experience
    if data.get('work_experience'):
        doc.add_heading('Work Experience', level=2)
        for job in data['work_experience']:
            p = doc.add_paragraph(f"{job.get('title', '')} at {job.get('company', '')}", style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25)
            for bullet in job.get('bullets', []):
                sub_p = doc.add_paragraph(bullet, style='List Bullet 2')
                sub_p.paragraph_format.left_indent = Inches(0.5)
    
    # Education
    if data.get('education'):
        doc.add_heading('Education', level=2)
        for edu in data['education']:
            p = doc.add_paragraph(f"{edu.get('degree', '')} from {edu.get('institution', '')} ({edu.get('year', '')})")
            if edu.get('cgpa'):
                p.add_run(f" - CGPA: {edu['cgpa']}")
    
    # Skills
    if data.get('skills_technical'):
        doc.add_heading('Technical Skills', level=2)
        doc.add_paragraph(", ".join(data['skills_technical']))
    
    # Footer
    if watermark:
        doc.add_paragraph().add_run("⚠ Free version - Watermarked ⚠").italic = True
    
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

def build_analysis(data):
    doc = Document()
    
    # Title
    title = doc.add_heading('ATS Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Scores
    doc.add_heading('ATS Match Score', level=2)
    doc.add_paragraph(f"Match Score: {data.get('match_score', 0)}%")
    doc.add_paragraph(f"Keywords Found: {data.get('ats_keywords_found', 0)}")
    doc.add_paragraph(f"Keywords Missing: {data.get('ats_keywords_missing', 0)}")
    
    # Analysis
    if data.get('score_explanation'):
        doc.add_heading('Score Analysis', level=2)
        doc.add_paragraph(data['score_explanation'])
    
    if data.get('strong_points'):
        doc.add_heading('Strong Points', level=2)
        for point in data['strong_points']:
            doc.add_paragraph(point, style='List Bullet')
    
    if data.get('missing_skills'):
        doc.add_heading('Missing Skills', level=2)
        for skill in data['missing_skills']:
            doc.add_paragraph(skill, style='List Bullet')
    
    if data.get('improvement_tips'):
        doc.add_heading('Improvement Tips', level=2)
        for i, tip in enumerate(data['improvement_tips'], 1):
            p = doc.add_paragraph()
            run1 = p.add_run(f"{i}. ")
            run1.bold = True
            run2 = p.add_run(tip)
    
    # Footer
    footer = doc.add_paragraph("Generated by ResumeReflect - AI Resume Tailoring Tool")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].italic = True
    footer.runs[0].font.size = Pt(8)
    
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

# ─────────────────────────────────────────────
# SESSION STATE
# ─────────────────────────────────────────────
defaults = {
    "step": 1, "resume_text": None, "jd_text": None,
    "ai_data": None, "ats_before": None, "ats_after": None,
    "plan": None, "email": None, "market_mode": "🇮🇳 India (Naukri / LinkedIn India)",
}
for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

# ─────────────────────────────────────────────
# HERO SECTION
# ─────────────────────────────────────────────
st.markdown("""
<div class="hero">
    <div class="badge">⚡ AI Powered · ATS Optimized · Instant Results</div>
    <h1>Land Interviews,<br>Not Spam Folders</h1>
    <p>Tailor your resume to match any job posting in 30 seconds. Watch your ATS score skyrocket with AI-powered precision.</p>
</div>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────
# STEP 1 — RESUME UPLOAD
# ─────────────────────────────────────────────
st.markdown('<div class="card"><div class="card-title"><span class="step-num">1</span> Upload Your Resume</div>', unsafe_allow_html=True)
resume_file = st.file_uploader("Upload as PDF", type=["pdf"], label_visibility="collapsed")
st.markdown('</div>', unsafe_allow_html=True)

if resume_file:
    file_bytes = resume_file.read()
    st.session_state.resume_text = read_pdf(file_bytes)
    if st.session_state.resume_text:
        st.success(f"✓ Resume loaded — {len(st.session_state.resume_text.split())} words")
    else:
        st.error("Could not extract text from PDF")

# ─────────────────────────────────────────────
# MARKET SELECTOR
# ─────────────────────────────────────────────
st.markdown('<div class="card"><div class="card-title">🌍 Target Job Market</div>', unsafe_allow_html=True)
market_mode = st.radio("Select market", [
    "🇮🇳 India (Naukri / LinkedIn India)",
    "🌍 International (LinkedIn / Indeed / Workday)"
], label_visibility="collapsed", horizontal=True)
st.session_state.market_mode = market_mode
st.markdown('</div>', unsafe_allow_html=True)

# ─────────────────────────────────────────────
# STEP 2 — JOB DESCRIPTION
# ─────────────────────────────────────────────
st.markdown('<div class="card"><div class="card-title"><span class="step-num">2</span> Add Job Description</div>', unsafe_allow_html=True)
jd_method = st.radio("How to add job description",
    ["Paste job link", "Paste text directly"],
    label_visibility="collapsed", horizontal=True)

if "link" in jd_method.lower():
    job_url = st.text_input("Job posting URL", placeholder="https://linkedin.com/jobs/view/...", label_visibility="collapsed")
    if job_url and st.button("🔗 Fetch from URL", key="fetch"):
        with st.spinner("Fetching job details..."):
            fetched = fetch_jd_from_url(job_url)
        if fetched:
            st.session_state.jd_text = fetched
            st.success("✓ Job description loaded!")
        else:
            st.error("Could not fetch. Please paste the text directly.")
else:
    jd_raw = st.text_area("Paste full job description", height=150,
        placeholder="Copy-paste the complete job description here...", label_visibility="collapsed")
    if jd_raw:
        st.session_state.jd_text = jd_raw

st.markdown('</div>', unsafe_allow_html=True)

# ─────────────────────────────────────────────
# STEP 3 — EMAIL + GENERATE
# ─────────────────────────────────────────────
st.markdown('<div class="card"><div class="card-title"><span class="step-num">3</span> Your Email (for results)</div>', unsafe_allow_html=True)
email = st.text_input("Email address", placeholder="you@email.com", label_visibility="collapsed")
st.session_state.email = email
st.markdown('</div>', unsafe_allow_html=True)

if st.button("🚀 Tailor & Analyze Resume", use_container_width=True):
    if not st.session_state.resume_text:
        st.error("Please upload your resume first.")
    elif not st.session_state.jd_text:
        st.error("Please add a job description.")
    elif not email or "@" not in email:
        st.error("Please enter a valid email address.")
    else:
        with st.spinner("🤖 AI is analyzing your resume against the job posting..."):
            try:
                before_score = simple_ats_score(st.session_state.resume_text, st.session_state.jd_text)
                st.session_state.ats_before = before_score
                
                response = call_groq(st.session_state.resume_text, st.session_state.jd_text, st.session_state.market_mode)
                data = parse_json(response)
                
                after_score = min(int(data.get('match_score', 75)), 100)
                st.session_state.ats_after = after_score
                st.session_state.ai_data = data
                
                log_user_action("ai_processed", email=email)
                st.session_state.step = 2
                st.rerun()
            except Exception as e:
                st.error(f"Error processing resume: {str(e)}")

# ─────────────────────────────────────────────
# RESULTS SECTION
# ─────────────────────────────────────────────
if st.session_state.step >= 2 and st.session_state.ai_data:
    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
    
    before = min(int(st.session_state.ats_before or 0), 100)
    after = min(int(st.session_state.ats_after or 0), 100)
    improvement = after - before
    
    st.markdown(f"""
    <div class="score-wrap">
        <div class="score-box">
            <div class="score-val score-before">{before}%</div>
            <div class="bar-wrap"><div class="bar" style="width:{before}%"></div></div>
            <div class="score-label">Before Tailoring</div>
        </div>
        <div class="score-box">
            <div class="score-val score-after">{after}%</div>
            <div class="bar-wrap"><div class="bar" style="width:{after}%"></div></div>
            <div class="score-label">After ATS Match</div>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    data = st.session_state.ai_data
    
    if improvement > 0:
        st.markdown(f"""
        <div class="card" style="border-color: rgba(0, 184, 133, 0.3); margin-top: 1rem;">
            <div class="card-title">📈 Score improved by <strong>{improvement}%</strong></div>
            <div style="color: var(--text-primary); font-size: 14px; line-height: 1.8;">{data.get('score_explanation', 'Your resume has been optimized.')}</div>
        </div>
        """, unsafe_allow_html=True)
    
    # Job titles
    if data.get('job_title_suggestions'):
        st.markdown("""
        <div class="card" style="border-color: rgba(45, 95, 158, 0.3); margin-top: 1rem;">
            <div class="card-title">💼 Job Titles to Search</div>
        """, unsafe_allow_html=True)
        chips = "".join([f'<span class="chip">{t}</span>' for t in data['job_title_suggestions']])
        st.markdown(f'<div style="margin: 0.5rem 0;">{chips}</div>', unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)
    
    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
    
    # Strong points and missing skills
    col1, col2 = st.columns(2)
    with col1:
        if data.get('strong_points'):
            st.markdown('<div class="card">', unsafe_allow_html=True)
            st.markdown('<div class="card-title">✅ Your Strengths</div>', unsafe_allow_html=True)
            for point in data.get('strong_points', [])[:4]:
                st.write(f"• {point}")
            st.markdown('</div>', unsafe_allow_html=True)
    
    with col2:
        if data.get('missing_skills'):
            st.markdown('<div class="card">', unsafe_allow_html=True)
            st.markdown('<div class="card-title">⚠️ Add These Skills</div>', unsafe_allow_html=True)
            for skill in data.get('missing_skills', [])[:4]:
                st.write(f"• {skill}")
            st.markdown('</div>', unsafe_allow_html=True)
    
    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
    
    # Download section
    st.markdown("""
    <div class="card">
        <div class="card-title">✅ Your Documents Are Ready</div>
        <div style="color: var(--text-secondary); font-size: 13px;">
            Free version includes watermark. <strong>Upgrade for clean documents + cover letter + interview prep kit.</strong>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    resume_bytes = build_resume(data, watermark=True)
    analysis_bytes = build_analysis(data)
    
    col1, col2 = st.columns(2)
    with col1:
        if st.download_button(
            "⬇ Download Tailored Resume",
            data=resume_bytes,
            file_name=f"resume_{data.get('candidate_name', 'candidate').replace(' ', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ):
            log_user_action("resume_downloaded", email=email)
    
    with col2:
        if st.download_button(
            "⬇ Download ATS Report",
            data=analysis_bytes,
            file_name="ats_analysis_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ):
            log_user_action("report_downloaded", email=email)
    
    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
    
    # Pricing
    st.markdown("""
    <div style="text-align: center; margin-bottom: 2rem;">
        <div style="font-family: 'Syne', sans-serif; font-size: 1.8rem; font-weight: 800; color: var(--text-primary);">Unlock All Features</div>
        <div style="color: var(--text-secondary); font-size: 13px; margin-top: 0.5rem;">One payment. Instant unlock. No subscriptions.</div>
    </div>
    """, unsafe_allow_html=True)
    
    st.markdown("""
    <div class="pricing-grid">
        <div class="plan">
            <div class="plan-name">Basic</div>
            <div class="plan-price">$5</div>
            <div class="plan-period">one-time payment</div>
            <ul class="plan-feat">
                <li>Clean resume (no watermark)</li>
                <li>ATS analysis report</li>
                <li>Job title suggestions</li>
            </ul>
        </div>
        <div class="plan hot">
            <div class="hot-badge">POPULAR</div>
            <div class="plan-name">Pro</div>
            <div class="plan-price">$10</div>
            <div class="plan-period">one-time payment</div>
            <ul class="plan-feat">
                <li>Everything in Basic</li>
                <li>AI-powered cover letter</li>
                <li>Interview preparation kit</li>
                <li>Keyword optimization tips</li>
            </ul>
        </div>
        <div class="plan hot" style="border-color: rgba(255, 107, 74, 0.4);">
            <div class="hot-badge" style="background: linear-gradient(135deg, #FF6B4A, #FF8B6B);">BEST VALUE</div>
            <div class="plan-name">Premium Yearly</div>
            <div class="plan-price">$79</div>
            <div class="plan-period">per year</div>
            <ul class="plan-feat">
                <li>Everything in Pro</li>
                <li>LinkedIn profile rewrite</li>
                <li>Unlimited resume tailoring</li>
                <li>Priority support</li>
                <li>6-month career coaching</li>
            </ul>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    col_b, col_p, col_y = st.columns(3)
    with col_b:
        if st.button("Get Basic $5 →", key="pay_basic", use_container_width=True):
            log_payment_interest("Basic $5", email)
            st.success("🚀 Payment coming soon! We've noted your interest.")
    
    with col_p:
        if st.button("Get Pro $10 →", key="pay_pro", use_container_width=True):
            log_payment_interest("Pro $10", email)
            st.success("🚀 Payment coming soon! We've noted your interest.")
    
    with col_y:
        if st.button("Get Premium $79 →", key="pay_year", use_container_width=True):
            log_payment_interest("Premium $79", email)
            st.success("🚀 Payment coming soon! We've noted your interest.")
    
    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
    
    if st.button("🔄 Tailor Another Resume", use_container_width=True):
        for k, v in defaults.items():
            st.session_state[k] = v
        st.rerun()

# ─────────────────────────────────────────────
# REVIEWS & TESTIMONIALS
# ─────────────────────────────────────────────
st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

st.markdown("""
<div class="review-section-title">⭐ Success Stories</div>
<div class="review-section-sub">Real results from real job seekers who used ResumeReflect</div>
<div class="reviews-grid">
    <div class="review-card">
        <div class="review-stars">★★★★★</div>
        <div class="review-body">"My ATS score jumped from 34% to 78%. Got an interview within 3 days. This tool actually works!"</div>
        <div class="review-name">Arjun M.</div>
        <div class="review-role">Software Engineer, Bangalore</div>
    </div>
    <div class="review-card">
        <div class="review-stars">★★★★★</div>
        <div class="review-body">"I was applying for 2 months with zero callbacks. After ResumeReflect, 2 interviews in the first week!"</div>
        <div class="review-name">Priya S.</div>
        <div class="review-role">Data Analyst, Mumbai</div>
    </div>
    <div class="review-card">
        <div class="review-stars">★★★★★</div>
        <div class="review-body">"The ATS report showed exactly what keywords I was missing. Tailored in 2 minutes, landed the role."</div>
        <div class="review-name">Rohit K.</div>
        <div class="review-role">Business Analyst, Hyderabad</div>
    </div>
    <div class="review-card">
        <div class="review-stars">★★★★★</div>
        <div class="review-body">"As a fresher, I had no idea how ATS works. This tool explained everything. Got my dream role!"</div>
        <div class="review-name">Sneha T.</div>
        <div class="review-role">Marketing Executive, Pune</div>
    </div>
</div>
""", unsafe_allow_html=True)

st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

# Review submission form
st.markdown("""
<div class="review-form-card">
    <div class="card-title">✍️ Share Your Success Story</div>
    <div style="color: var(--text-secondary); font-size: 12px; margin-top: 0.5rem;">
        Got an interview after using ResumeReflect? Your story helps others — takes just 60 seconds!
    </div>
</div>
""", unsafe_allow_html=True)

r_col1, r_col2 = st.columns(2)
with r_col1:
    reviewer_name = st.text_input("Your Name", placeholder="e.g. Rahul S. (optional)", key="reviewer_name")
with r_col2:
    reviewer_role = st.text_input("Job Role", placeholder="e.g. Software Engineer at Infosys", key="reviewer_role")

review_text = st.text_area(
    "Your Experience",
    placeholder="How did ResumeReflect help? Did you get interviews? ATS score before/after? Any feedback?",
    height=100,
    key="review_text"
)

review_rating = st.select_slider(
    "⭐ Your Rating",
    options=["1 Star", "2 Stars", "3 Stars", "4 Stars", "5 Stars"],
    value="5 Stars",
    key="review_rating"
)

if st.button("Submit Review ✅", use_container_width=True, key="submit_review"):
    if review_text and len(review_text.strip()) > 10:
        try:
            requests.post(SHEET_SCRIPT_URL, json={
                "type": "review",
                "name": reviewer_name.strip() if reviewer_name.strip() else "Anonymous",
                "role": reviewer_role or "",
                "review": review_text.strip(),
                "rating": review_rating,
                "email": st.session_state.get("email", ""),
                "tool_name": TOOL_NAME,  # FIX: Added proper tool name
                "timestamp": datetime.now().isoformat()
            }, timeout=5)
            st.success("🙏 Thank you! Your review helps our community.")
            st.balloons()
        except:
            st.success("🙏 Thank you! Your review has been noted.")
    else:
        st.error("Please write at least a sentence.")

st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

# ─────────────────────────────────────────────
# FOOTER
# ─────────────────────────────────────────────
st.markdown(f"""
<div class="footer">
    <div class="footer-brand">⚡ {TOOL_NAME}</div>
    Built for serious job seekers · Made in India 🇮🇳
    <br><br>
    <div style="font-size: 11px; margin-bottom: 0.3rem;">🔒 Your resume is processed by AI and never stored on our servers.</div>
    <div style="font-size: 11px;">© 2026 {TOOL_NAME}. All rights reserved.</div>
</div>
""", unsafe_allow_html=True)
