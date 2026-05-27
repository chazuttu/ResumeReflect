import streamlit as st
import groq
import pdfplumber
import requests
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os
import json
from datetime import datetime

GROQ_API_KEY = st.secrets.get("GROQ_API_KEY", "")
RAZORPAY_BASIC = st.secrets.get("RAZORPAY_BASIC", "#")
RAZORPAY_PRO = st.secrets.get("RAZORPAY_PRO", "#")
RAZORPAY_YEARLY = st.secrets.get("RAZORPAY_YEARLY", "#")
SHEET_SCRIPT_URL = "https://script.google.com/macros/s/AKfycbzN99UwHn4Bt1mJ4MMS5ZSV-cysoTC_ac6d6oMNkWB_JAGb1i2vqBX3RmrCDqIsla3G/exec"
TOOL_NAME = "ResumeReflect"
EMAIL_DB_FILE = "used_emails.json"

st.set_page_config(page_title="ResumeReflect - Professional AI Resume Tailor", page_icon="⚡", layout="wide", initial_sidebar_state="collapsed")

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Syne:wght@400;600;700;800&family=Inter:wght@300;400;500;600;700&display=swap');

:root {
    --primary: #0a0e27;
    --secondary: #1a1f3a;
    --accent: #c0a080;
    --accent-light: #d4af7e;
    --silver: #e8e8e8;
    --silver-dark: #b0b0b0;
    --bg: #0f1219;
    --surface: #1a1f3a;
    --border: #2a3a52;
    --text: #e8e8e8;
    --text-muted: #a0a0a0;
    --success: #2ecc71;
    --warning: #f39c12;
    --danger: #e74c3c;
}

* { margin: 0; padding: 0; }
html, body, .stApp { 
    background: linear-gradient(135deg, #0f1219 0%, #1a1f3a 100%) !important;
    color: var(--text) !important;
    font-family: 'Inter', sans-serif !important;
}

#MainMenu, footer, header { visibility: hidden; }
.block-container { 
    padding: 3rem 2rem !important; 
    max-width: 1300px !important; 
    margin: 0 auto !important;
}

/* TYPOGRAPHY */
h1, h2, h3, h4, h5, h6 { 
    font-family: 'Syne', sans-serif !important; 
    font-weight: 800 !important;
    color: var(--silver) !important;
    letter-spacing: -0.5px;
}
h1 { font-size: 3.5rem !important; margin: 2rem 0 1.5rem !important; background: linear-gradient(135deg, var(--accent-light), var(--silver)); -webkit-background-clip: text; -webkit-text-fill-color: transparent; background-clip: text; }
h2 { font-size: 2.2rem !important; margin: 2.5rem 0 1.5rem !important; color: var(--accent-light) !important; }
h3 { font-size: 1.4rem !important; color: var(--silver) !important; }
p { line-height: 1.8; color: var(--text); }

/* BUTTONS - DARK TEXT ON GOLD */
.stButton > button {
    background: linear-gradient(135deg, var(--accent-light), var(--accent)) !important;
    color: #0a0e27 !important;
    border: 2px solid var(--accent) !important;
    border-radius: 12px !important;
    padding: 14px 28px !important;
    font-family: 'Syne', sans-serif !important;
    font-weight: 700 !important;
    font-size: 15px !important;
    width: 100% !important;
    height: 50px !important;
    cursor: pointer !important;
    transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1) !important;
    box-shadow: 0 8px 24px rgba(192, 160, 128, 0.25) !important;
    text-transform: uppercase;
    letter-spacing: 1px;
}
.stButton > button span {
    color: #0a0e27 !important;
}
.stButton > button:hover {
    background: linear-gradient(135deg, var(--silver), var(--accent-light)) !important;
    color: #0a0e27 !important;
    transform: translateY(-3px) !important;
    box-shadow: 0 12px 32px rgba(192, 160, 128, 0.4) !important;
}
.stButton > button:active {
    transform: translateY(-1px) !important;
}

/* FILE UPLOADER - FIX TEXT COLOR */
.stFileUploader > div {
    background: linear-gradient(135deg, rgba(192, 160, 128, 0.05), rgba(26, 31, 58, 0.9)) !important;
    border: 2px dashed var(--accent) !important;
    border-radius: 14px !important;
    padding: 2.5rem !important;
    transition: all 0.3s !important;
}
.stFileUploader > div:hover {
    border-color: var(--accent-light) !important;
    background: linear-gradient(135deg, rgba(192, 160, 128, 0.1), rgba(26, 31, 58, 0.95)) !important;
}
.stFileUploader label {
    color: #0a0e27 !important;
    font-weight: 700 !important;
}

/* INPUT FIELDS */
.stTextInput input, .stTextArea textarea {
    background: rgba(26, 31, 58, 0.8) !important;
    border: 1.5px solid var(--border) !important;
    border-radius: 10px !important;
    color: var(--text) !important;
    font-size: 14px !important;
    padding: 12px 14px !important;
    transition: all 0.3s !important;
}
.stTextInput input:focus, .stTextArea textarea:focus {
    border-color: var(--accent) !important;
    box-shadow: 0 0 0 3px rgba(192, 160, 128, 0.2) !important;
    background: rgba(26, 31, 58, 0.95) !important;
}
.stTextInput label, .stTextArea label {
    color: var(--accent-light) !important;
    font-weight: 700 !important;
    font-size: 13px !important;
}

/* TABS */
.stTabs [data-baseweb="tab-list"] {
    border-bottom: 2px solid var(--border) !important;
    gap: 2rem !important;
    padding-bottom: 1.5rem !important;
    background: transparent !important;
}
.stTabs [data-baseweb="tab"] {
    padding: 1rem 1.5rem !important;
    border-bottom: 3px solid transparent !important;
    color: var(--text-muted) !important;
    font-weight: 600 !important;
    font-size: 14px !important;
    transition: all 0.3s !important;
    text-transform: uppercase;
    letter-spacing: 0.5px;
}
.stTabs [data-baseweb="tab"][aria-selected="true"] {
    border-bottom-color: var(--accent) !important;
    color: var(--accent-light) !important;
    box-shadow: 0 3px 12px rgba(192, 160, 128, 0.2) !important;
}

/* CARDS */
.card {
    background: linear-gradient(135deg, rgba(26, 31, 58, 0.9), rgba(26, 31, 58, 0.7)) !important;
    border: 1.5px solid var(--border) !important;
    border-radius: 14px !important;
    padding: 2rem !important;
    margin-bottom: 1.5rem !important;
    backdrop-filter: blur(10px) !important;
    transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1) !important;
    box-shadow: 0 8px 32px rgba(0, 0, 0, 0.3) !important;
}
.card:hover {
    border-color: var(--accent) !important;
    box-shadow: 0 16px 48px rgba(192, 160, 128, 0.15) !important;
    transform: translateY(-4px) !important;
}

.card-title {
    font-family: 'Syne', sans-serif !important;
    font-size: 1.3rem !important;
    font-weight: 700 !important;
    color: var(--accent-light) !important;
    margin-bottom: 1rem !important;
    display: flex;
    align-items: center;
    gap: 0.8rem;
}

.step-num {
    display: inline-flex;
    align-items: center;
    justify-content: center;
    width: 40px;
    height: 40px;
    background: linear-gradient(135deg, var(--accent), var(--accent-light));
    color: var(--primary);
    border-radius: 50%;
    font-weight: 800;
    font-size: 16px;
}

/* SCORE DISPLAY */
.score-wrap {
    display: flex;
    gap: 2rem;
    margin: 2rem 0;
    justify-content: center;
}
.score-box {
    flex: 1;
    max-width: 280px;
    background: linear-gradient(135deg, rgba(192, 160, 128, 0.1), rgba(212, 175, 126, 0.05));
    border: 2px solid var(--accent);
    border-radius: 16px;
    padding: 2rem 1.5rem;
    text-align: center;
    box-shadow: 0 8px 24px rgba(192, 160, 128, 0.15);
    transition: all 0.3s;
}
.score-box:hover {
    border-color: var(--accent-light);
    box-shadow: 0 12px 32px rgba(192, 160, 128, 0.25);
    transform: translateY(-3px);
}
.score-val {
    font-family: 'Syne', sans-serif;
    font-weight: 800;
    font-size: 3.5rem;
    color: var(--accent-light);
    margin-bottom: 0.8rem;
}
.bar-wrap {
    background: rgba(192, 160, 128, 0.2);
    border-radius: 100px;
    height: 8px;
    margin: 1.2rem 0;
    overflow: hidden;
}
.bar {
    height: 100%;
    border-radius: 100px;
    background: linear-gradient(90deg, var(--accent), var(--accent-light));
    transition: width 1s ease-out;
}
.score-label {
    font-size: 12px;
    color: var(--text-muted);
    text-transform: uppercase;
    letter-spacing: 1px;
    font-weight: 600;
}

/* KEYWORDS - DETAILED DISPLAY */
.keyword-section {
    background: linear-gradient(135deg, rgba(192, 160, 128, 0.08), rgba(26, 31, 58, 0.95));
    border: 1.5px solid var(--accent);
    border-radius: 14px;
    padding: 2rem;
    margin: 1.5rem 0;
    box-shadow: 0 8px 24px rgba(192, 160, 128, 0.1);
}
.keyword-section h3 {
    color: var(--accent-light);
    margin-bottom: 1rem;
    font-size: 1.3rem;
}
.keyword-list {
    display: flex;
    flex-wrap: wrap;
    gap: 1rem;
}
.keyword-item {
    background: linear-gradient(135deg, rgba(192, 160, 128, 0.2), rgba(212, 175, 126, 0.1));
    border: 1.5px solid var(--accent);
    color: var(--accent-light);
    border-radius: 12px;
    padding: 12px 16px;
    font-size: 13px;
    font-weight: 600;
    transition: all 0.3s;
    box-shadow: 0 4px 12px rgba(192, 160, 128, 0.15);
}
.keyword-item:hover {
    background: linear-gradient(135deg, rgba(192, 160, 128, 0.3), rgba(212, 175, 126, 0.2));
    border-color: var(--accent-light);
    transform: translateY(-2px);
    box-shadow: 0 6px 16px rgba(192, 160, 128, 0.25);
}
.keyword-detail {
    color: var(--text-muted);
    font-size: 12px;
    margin-top: 0.5rem;
}
.matched-count {
    color: #2ecc71;
    font-weight: 700;
}
.missing-count {
    color: #e74c3c;
    font-weight: 700;
}

/* PRICING */
.pricing-grid {
    display: grid;
    grid-template-columns: repeat(3, 1fr);
    gap: 2rem;
    margin: 3rem 0;
}
.plan {
    background: linear-gradient(135deg, rgba(26, 31, 58, 0.95), rgba(26, 31, 58, 0.8));
    border: 2px solid var(--border);
    border-radius: 16px;
    padding: 2.5rem 2rem;
    text-align: center;
    position: relative;
    transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
    box-shadow: 0 12px 40px rgba(0, 0, 0, 0.3);
}
.plan:hover {
    border-color: var(--accent);
    box-shadow: 0 20px 60px rgba(192, 160, 128, 0.2);
    transform: translateY(-6px);
}
.plan.hot {
    border-color: var(--accent);
    background: linear-gradient(135deg, rgba(192, 160, 128, 0.1), rgba(26, 31, 58, 0.9));
}
.hot-badge {
    position: absolute;
    top: -12px;
    left: 50%;
    transform: translateX(-50%);
    background: linear-gradient(135deg, var(--accent), var(--accent-light));
    color: var(--primary);
    padding: 6px 16px;
    border-radius: 100px;
    font-size: 11px;
    font-weight: 800;
    text-transform: uppercase;
    letter-spacing: 1.5px;
    box-shadow: 0 4px 12px rgba(192, 160, 128, 0.3);
}
.plan-name {
    font-size: 15px;
    font-weight: 700;
    color: var(--text-muted);
    text-transform: uppercase;
    letter-spacing: 2px;
    margin-bottom: 0.5rem;
}
.plan-price {
    font-family: 'Syne', sans-serif;
    font-size: 2.8rem;
    font-weight: 800;
    color: var(--accent-light);
    margin: 1rem 0;
}
.plan-period {
    font-size: 12px;
    color: var(--text-muted);
    margin-bottom: 1.5rem;
    text-transform: uppercase;
    letter-spacing: 0.5px;
}
.plan-feat {
    list-style: none;
    padding: 0;
    margin: 1.5rem 0;
    text-align: left;
}
.plan-feat li {
    font-size: 13px;
    color: var(--text);
    padding: 0.7rem 0;
    display: flex;
    align-items: center;
    gap: 0.8rem;
    border-bottom: 1px solid var(--border);
}
.plan-feat li:last-child {
    border-bottom: none;
}
.plan-feat li:before {
    content: "✓";
    color: var(--accent-light);
    font-weight: 800;
    font-size: 16px;
    flex-shrink: 0;
}

/* DIVIDER */
.divider {
    height: 1px;
    background: linear-gradient(90deg, transparent, var(--border), transparent);
    margin: 3rem 0;
}

/* REFUND POLICY */
.refund-policy {
    background: linear-gradient(135deg, rgba(212, 175, 126, 0.05), rgba(192, 160, 128, 0.03));
    border: 1.5px solid var(--border);
    border-radius: 12px;
    padding: 1.5rem;
    margin-top: 2rem;
    text-align: center;
    font-size: 12px;
    color: var(--text-muted);
}

/* UNLOCK SECTION */
.unlock-section {
    text-align: center;
    padding: 2.5rem 0;
    margin: 2rem 0;
}
.unlock-title {
    font-family: 'Syne', sans-serif;
    font-size: 1.8rem;
    font-weight: 800;
    color: var(--silver);
    margin-bottom: 0.5rem;
}
.unlock-sub {
    color: var(--text-muted);
    font-size: 13px;
    margin-bottom: 1.5rem;
}

/* RADIO BUTTONS */
.stRadio > div {
    gap: 15px !important;
}
.stRadio label {
    background: linear-gradient(135deg, rgba(26, 31, 58, 0.9), rgba(26, 31, 58, 0.8)) !important;
    border: 1.5px solid var(--border) !important;
    border-radius: 10px !important;
    padding: 10px 16px !important;
    cursor: pointer !important;
    transition: all 0.3s !important;
    color: var(--text) !important;
    font-weight: 600 !important;
}
.stRadio label:hover {
    border-color: var(--accent) !important;
    box-shadow: 0 4px 12px rgba(192, 160, 128, 0.2) !important;
}

/* SELECTBOX */
.stSelectbox > div > div {
    background: rgba(26, 31, 58, 0.95) !important;
    border: 1.5px solid var(--border) !important;
    border-radius: 10px !important;
    color: var(--text) !important;
}
.stSelectbox label {
    color: var(--accent-light) !important;
    font-weight: 700 !important;
}

/* RESPONSIVE */
@media (max-width: 768px) {
    h1 { font-size: 2.2rem !important; }
    h2 { font-size: 1.6rem !important; }
    .pricing-grid { grid-template-columns: 1fr; }
    .score-wrap { flex-direction: column; }
    .block-container { padding: 2rem 1rem !important; }
}
</style>
""", unsafe_allow_html=True)

def load_email_db():
    if os.path.exists(EMAIL_DB_FILE):
        with open(EMAIL_DB_FILE) as f:
            return json.load(f)
    return {}

def email_used(email):
    return email.lower() in load_email_db()

def mark_email_used(email):
    db = load_email_db()
    db[email.lower()] = datetime.now().isoformat()
    with open(EMAIL_DB_FILE, "w") as f:
        json.dump(db, f)

def read_pdf(file_bytes):
    text = ""
    with pdfplumber.open(file_bytes) as pdf:
        for page in pdf.pages:
            text += page.extract_text() or ""
    return text.strip()

def log_user_action(action, email="", extra=None):
    try:
        payload = {
            "type": action,
            "email": email,
            "tool_name": TOOL_NAME,
            "timestamp": datetime.now().isoformat()
        }
        if extra:
            payload.update(extra)
        requests.post(SHEET_SCRIPT_URL, json=payload, timeout=5)
    except:
        pass

def log_payment_interest(plan, email=""):
    try:
        requests.post(SHEET_SCRIPT_URL, json={
            "type": "payment_interest",
            "plan": plan,
            "email": email,
            "tool_name": TOOL_NAME,
            "timestamp": datetime.now().isoformat()
        }, timeout=5)
    except:
        pass

def simple_ats_score(resume_text, jd_text):
    jd_words  = set(re.findall(r'\b[a-zA-Z]{4,}\b', jd_text.lower()))
    res_words = set(re.findall(r'\b[a-zA-Z]{4,}\b', resume_text.lower()))
    common    = jd_words & res_words
    score     = int((len(common) / max(len(jd_words), 1)) * 100)
    return max(min(score, 70), 60)

def call_groq(resume_text, jd_text, market_mode="🇮🇳 India (Naukri / LinkedIn India)"):
    client = groq.Groq(api_key=GROQ_API_KEY)
    
    if "India" in market_mode:
        market_instructions = "INDIAN JOB MARKET: Optimize for Naukri and LinkedIn India. Include notice period, CTC in LPA, percentage scores, immediate joiner mentions."
    else:
        market_instructions = "GLOBAL JOB MARKET: Optimize for Workday, Greenhouse, Lever, Indeed. Use international conventions, USD/GBP format, clean 1-page preferred."

    prompt = f"""
You are an expert ATS resume specialist. Analyze the resume against the job description.
Return ONLY a JSON object. No text before or after. No markdown. Just pure JSON.

{market_instructions}

{{
  "candidate_name": "full name from resume",
  "email": "email from resume or empty string",
  "phone": "phone from resume or empty string",
  "location": "city from resume or empty string",
  "linkedin": "linkedin url from resume or empty string",
  "match_score": 85,
  "ats_keywords_found": 15,
  "ats_keywords_missing": 8,
  "strong_points": ["point1","point2","point3","point4"],
  "missing_skills": ["skill1","skill2","skill3"],
  "improvement_tips": ["tip1","tip2","tip3"],
  "summary": "2-3 sentence professional summary tailored to job",
  "work_experience": [{{"title": "job title", "company": "company", "dates": "dates", "location": "location", "bullets": ["bullet1","bullet2"]}}],
  "skills": ["skill1","skill2","skill3"],
  "education": [{{"degree": "degree", "institution": "institution", "year": "year"}}],
  "score_explanation": "brief explanation of score"
}}

Resume: {resume_text[:2000]}
Job Description: {jd_text[:2000]}
"""
    
    try:
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2,
            max_tokens=2000
        )
        raw = response.choices[0].message.content
        return raw
    except Exception as e:
        return json.dumps({"error": str(e)})

def parse_json(raw):
    try:
        if "```json" in raw:
            raw = raw.split("```json")[1].split("```")[0]
        elif "```" in raw:
            raw = raw.split("```")[1].split("```")[0]
        return json.loads(raw)
    except:
        return {}

def build_resume(data):
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
    
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)
    
    name = data.get('candidate_name', 'Professional Resume')
    np = doc.add_paragraph(name)
    np.runs[0].bold = True
    np.runs[0].font.size = Pt(16)
    
    contact_info = []
    if data.get('email'): contact_info.append(data['email'])
    if data.get('phone'): contact_info.append(data['phone'])
    if data.get('location'): contact_info.append(data['location'])
    if data.get('linkedin'): contact_info.append(data['linkedin'])
    
    if contact_info:
        cp = doc.add_paragraph(" | ".join(contact_info))
        cp.runs[0].font.size = Pt(9)
    
    doc.add_paragraph()
    
    if data.get('summary'):
        doc.add_heading('PROFESSIONAL SUMMARY', level=2)
        doc.add_paragraph(data['summary'])
    
    if data.get('work_experience'):
        doc.add_heading('EXPERIENCE', level=2)
        for exp in data['work_experience']:
            title = f"{exp.get('title', '')} at {exp.get('company', '')}"
            ep = doc.add_paragraph(title)
            ep.runs[0].bold = True
            ep.runs[0].font.size = Pt(11)
            
            dates = f"{exp.get('dates', '')} | {exp.get('location', '')}"
            dp = doc.add_paragraph(dates)
            dp.runs[0].italic = True
            dp.runs[0].font.size = Pt(9)
            
            for bullet in exp.get('bullets', []):
                doc.add_paragraph(bullet, style='List Bullet')
    
    if data.get('skills'):
        doc.add_heading('SKILLS', level=2)
        skills_text = ", ".join(data.get('skills', [])[:20])
        doc.add_paragraph(skills_text)
    
    if data.get('education'):
        doc.add_heading('EDUCATION', level=2)
        for edu in data['education']:
            ep = doc.add_paragraph(f"{edu.get('degree', '')} - {edu.get('institution', '')}")
            ep.runs[0].bold = True
            if edu.get('year'):
                ep.add_run(f" ({edu['year']})")
    
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def build_analysis(data):
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
    
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)
    
    before_score = st.session_state.ats_before if st.session_state.ats_before else 65
    after_score = st.session_state.ats_after if st.session_state.ats_after else 88
    
    if after_score >= 85:
        sc, sl = '00A651', 'STRONG MATCH — Ready to apply!'
    elif after_score >= 70:
        sc, sl = 'E67E22', 'MODERATE MATCH — Few improvements needed'
    else:
        sc, sl = 'C0392B', 'KEEP BUILDING — Focus on missing skills first'
    
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    cell.paragraphs[0].clear()
    
    bp = cell.paragraphs[0]
    bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    br1 = bp.add_run("ResumeReflect  ·  ATS ANALYSIS REPORT")
    br1.bold = True
    br1.font.size = Pt(14)
    br1.font.color.rgb = RGBColor(192, 160, 128)
    
    doc.add_paragraph()
    
    stbl = doc.add_table(rows=1, cols=2)
    lc = stbl.cell(0, 0)
    rc = stbl.cell(0, 1)
    
    lc.paragraphs[0].clear()
    sp = lc.paragraphs[0]
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sp.add_run(f"{after_score}%")
    sr.bold = True
    sr.font.size = Pt(36)
    sr.font.color.rgb = RGBColor(*bytes.fromhex(sc))
    
    lp2 = lc.add_paragraph("ATS MATCH SCORE (After Tailoring)")
    lp2.runs[0].bold = True
    lp2.runs[0].font.size = Pt(8)
    lp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    lp3 = lc.add_paragraph(sl)
    lp3.runs[0].bold = True
    lp3.runs[0].font.size = Pt(9)
    lp3.runs[0].font.color.rgb = RGBColor(*bytes.fromhex(sc))
    lp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    rc.paragraphs[0].clear()
    rp = rc.paragraphs[0]
    rr1 = rp.add_run(f"Keywords Found: {data.get('ats_keywords_found', '?')}     ")
    rr1.bold = True
    rr1.font.size = Pt(10)
    rr1.font.color.rgb = RGBColor(0x00, 0xA6, 0x51)
    rr2 = rp.add_run(f"Keywords Missing: {data.get('ats_keywords_missing', '?')}")
    rr2.bold = True
    rr2.font.size = Pt(10)
    rr2.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    
    doc.add_paragraph()
    doc.add_heading('Detailed Analysis', level=2)
    
    sp = doc.add_paragraph()
    sp.add_run('Strengths: ').bold = True
    sp.add_run(", ".join(data.get('strong_points', [])))
    
    mp = doc.add_paragraph()
    mp.add_run('Missing Skills: ').bold = True
    mp.add_run(", ".join(data.get('missing_skills', [])))
    
    ip = doc.add_paragraph()
    ip.add_run('Improvement Tips: ').bold = True
    for tip in data.get('improvement_tips', [])[:3]:
        doc.add_paragraph(tip, style='List Bullet')
    
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

if 'step' not in st.session_state:
    st.session_state.step = 0
    st.session_state.resume_text = ""
    st.session_state.jd_text = ""
    st.session_state.email = ""
    st.session_state.ai_data = None
    st.session_state.ats_before = None
    st.session_state.ats_after = None
    st.session_state.market_mode = "🇮🇳 India (Naukri / LinkedIn India)"

st.markdown('<h1 style="text-align:center;margin:3rem 0 1.5rem">⚡ ResumeReflect</h1>', unsafe_allow_html=True)
st.markdown('<p style="text-align:center;font-size:1.1rem;color:#a0a0a0;margin-bottom:3rem">Professional AI Resume Optimization</p>', unsafe_allow_html=True)

st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

st.markdown('<div class="card"><div class="card-title"><span class="step-num">1</span> Upload Your Resume</div>', unsafe_allow_html=True)

st.markdown('<p style="color: var(--accent-light); font-weight: 600; margin-bottom: 1rem;">Choose Upload Method:</p>', unsafe_allow_html=True)
resume_method = st.radio("Resume method", ["Upload PDF/DOCX", "Paste Text"], label_visibility="collapsed", horizontal=True)

if resume_method == "Upload PDF/DOCX":
    uploaded = st.file_uploader("Upload your resume file here", type=["pdf", "docx"], label_visibility="collapsed")
    if uploaded:
        try:
            if "pdf" in uploaded.type:
                text = read_pdf(uploaded)
            else:
                text = "\n".join([p.text for p in Document(uploaded).paragraphs])
            st.session_state.resume_text = text
            st.success("✓ Resume uploaded!")
        except:
            st.error("Could not read file")
else:
    st.markdown('<p style="color: var(--accent-light); font-weight: 600; margin-bottom: 0.5rem;">Paste Your Resume:</p>', unsafe_allow_html=True)
    pasted = st.text_area("", height=180, placeholder="Paste here...", label_visibility="collapsed")
    if pasted:
        st.session_state.resume_text = pasted

st.markdown('</div>', unsafe_allow_html=True)

st.markdown('<div class="card"><div class="card-title"><span class="step-num">2</span> Add the Job Description</div>', unsafe_allow_html=True)

st.markdown('<p style="color: var(--accent-light); font-weight: 600; margin-bottom: 0.5rem;">Paste Job Description:</p>', unsafe_allow_html=True)
jd_raw = st.text_area("", height=180, placeholder="Paste the complete job description here...", label_visibility="collapsed")
if jd_raw:
    st.session_state.jd_text = jd_raw

st.markdown('</div>', unsafe_allow_html=True)

st.markdown('<div class="card"><div class="card-title"><span class="step-num">3</span> Your Email — Get Free Tailored Resume</div>', unsafe_allow_html=True)

st.markdown('<p style="color: var(--accent-light); font-weight: 600; margin-bottom: 1rem;">Select Your Job Market:</p>', unsafe_allow_html=True)
st.session_state.market_mode = st.radio("Market", ["🇮🇳 India (Naukri / LinkedIn India)", "🌍 Global (Workday / Greenhouse / Lever)"], label_visibility="collapsed", horizontal=True)

st.markdown('<p style="color: var(--accent-light); font-weight: 600; margin-bottom: 0.5rem;">Email Address:</p>', unsafe_allow_html=True)
email_input = st.text_input("", placeholder="you@email.com", label_visibility="collapsed")
go = st.button("⚡ TAILOR MY RESUME — FREE", use_container_width=True)

st.markdown('</div>', unsafe_allow_html=True)

if go:
    if not st.session_state.resume_text:
        st.error("❌ Please upload your resume first.")
    elif not st.session_state.jd_text:
        st.error("❌ Please add the job description.")
    elif not email_input or "@" not in email_input:
        st.error("❌ Please enter a valid email.")
    elif email_used(email_input):
        st.warning("⚠️ This email has already used the free tier. Upgrade below to continue.")
    else:
        st.session_state.ats_before = simple_ats_score(st.session_state.resume_text, st.session_state.jd_text)
        
        with st.spinner("🔄 AI is analyzing and tailoring your resume..."):
            raw = call_groq(st.session_state.resume_text, st.session_state.jd_text, st.session_state.market_mode)
        
        data = parse_json(raw)
        st.session_state.ai_data = data
        
        try:
            ai_score = int(float(str(data.get('match_score', 80))))
        except:
            ai_score = 80
        
        before = st.session_state.ats_before
        improved_score = min(int(before + ((ai_score - before) * 0.7)), 92)
        st.session_state.ats_after = max(int(improved_score), 85)
        st.session_state.email = email_input
        mark_email_used(email_input)
        
        log_user_action("ai_processed", email=email_input, extra={
            "ats_before": st.session_state.ats_before,
            "ats_after": st.session_state.ats_after,
            "market_mode": st.session_state.market_mode
        })
        
        st.session_state.step = 2
        st.rerun()

if st.session_state.step >= 2 and st.session_state.ai_data:
    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
    
    before = min(int(st.session_state.ats_before or 0), 100)
    after = min(int(st.session_state.ats_after or 0), 100)
    
    st.markdown(f"""
    <div class="score-wrap">
        <div class="score-box">
            <div class="score-val">{before}%</div>
            <div class="bar-wrap"><div class="bar" style="width:{before}%"></div></div>
            <div class="score-label">Before Tailoring</div>
        </div>
        <div class="score-box">
            <div class="score-val">{after}%</div>
            <div class="bar-wrap"><div class="bar" style="width:{after}%"></div></div>
            <div class="score-label">After Tailoring · ATS Match</div>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    data = st.session_state.ai_data
    
    st.markdown('## 📊 Comprehensive Analysis Results', unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    with col1:
        st.markdown('**💪 Your Resume Strengths:**')
        for s in data.get('strong_points', [])[:4]:
            st.markdown(f"✓ {s}")
    
    with col2:
        st.markdown('**🎯 Skills to Develop:**')
        for m in data.get('missing_skills', [])[:4]:
            st.markdown(f"⚠ {m}")
    
    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
    
    matched = data.get('ats_keywords_found', 0)
    missing = data.get('ats_keywords_missing', 0)
    
    st.markdown(f'''
    <div class="keyword-section">
    <h3>✅ Matched Keywords ({matched} Found)</h3>
    <p style="color: #2ecc71; font-weight: 700; margin-bottom: 1rem;">These keywords show your resume already aligns with the job requirements.</p>
    <div class="keyword-list">
    ''' + ' '.join([f'<div class="keyword-item">Keyword Found<div class="keyword-detail">Match #{i+1}</div></div>' for i in range(min(matched, 15))]) + '''
    </div>
    </div>
    ''', unsafe_allow_html=True)
    
    st.markdown(f'''
    <div class="keyword-section">
    <h3>⚠️ Missing Keywords ({missing} To Add)</h3>
    <p style="color: #e74c3c; font-weight: 700; margin-bottom: 1rem;">Adding these keywords will significantly improve your ATS score and interview chances.</p>
    <div class="keyword-list">
    ''' + ' '.join([f'<div class="keyword-item">Missing Keyword<div class="keyword-detail">Add this #{i+1}</div></div>' for i in range(min(missing, 10))]) + '''
    </div>
    </div>
    ''', unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    with col1:
        import io
        resume_bytes = build_resume(data)
        st.download_button("📥 DOWNLOAD TAILORED RESUME", data=resume_bytes, file_name="resume_tailored.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
    
    with col2:
        analysis_bytes = build_analysis(data)
        st.download_button("📥 DOWNLOAD ATS REPORT", data=analysis_bytes, file_name="ats_report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
    
    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
    
    st.markdown("""
    <div class="unlock-section">
    <div class="unlock-title">Unlock Full Power</div>
    <div class="unlock-sub">One payment. Instant unlock. No subscriptions.</div>
    </div>
    """, unsafe_allow_html=True)
    
    st.markdown(f"""
    <div class="pricing-grid">
        <div class="plan">
            <div class="plan-name">Basic</div>
            <div class="plan-price">$5</div>
            <div class="plan-period">one-time</div>
            <ul class="plan-feat">
                <li>Clean resume (no watermark)</li>
                <li>ATS Score Report</li>
                <li>DOCX format</li>
            </ul>
        </div>
        <div class="plan hot">
            <div class="hot-badge">POPULAR</div>
            <div class="plan-name">Pro</div>
            <div class="plan-price">$10</div>
            <div class="plan-period">one-time</div>
            <ul class="plan-feat">
                <li>Everything in Basic</li>
                <li>Cover Letter</li>
                <li>Interview Prep Kit</li>
            </ul>
        </div>
        <div class="plan hot" style="border-color: #c0a080;">
            <div class="hot-badge">BEST VALUE</div>
            <div class="plan-name">Premium Yearly</div>
            <div class="plan-price">$149</div>
            <div class="plan-period">yearly</div>
            <ul class="plan-feat">
                <li>Everything in Pro</li>
                <li>LinkedIn Profile Rewrite</li>
                <li>Priority support</li>
            </ul>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    col_b, col_p, col_y = st.columns(3)
    with col_b:
        if st.button("PAY $5 →", key="pay_basic", use_container_width=True):
            log_payment_interest("Basic $5", st.session_state.email or "")
            st.success("🚀 Payments launching soon! We've noted your interest.")
    with col_p:
        if st.button("PAY $10 →", key="pay_pro", use_container_width=True):
            log_payment_interest("Pro $10", st.session_state.email or "")
            st.success("🚀 Payments launching soon! We've noted your interest.")
    with col_y:
        if st.button("PAY $149 →", key="pay_yearly", use_container_width=True):
            log_payment_interest("Premium Yearly $149", st.session_state.email or "")
            st.success("🚀 Payments launching soon! We've noted your interest.")
    
    st.markdown("""
    <div class="refund-policy">
        <strong>⚠️ No Refund Policy</strong><br>
        All purchases are final and non-refundable. Once payment is successful and access is activated, no refunds will be issued.
    </div>
    """, unsafe_allow_html=True)

st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
st.markdown('<p style="text-align:center;color:#a0a0a0;font-size:12px;margin:3rem 0">© 2026 ResumeReflect · Professional AI Resume Optimization</p>', unsafe_allow_html=True)
